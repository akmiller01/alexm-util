import csv
from docx import Document
import django
from django.template import Template, Context
from django.conf import settings
TEMPLATES = [
    {
        'BACKEND': 'django.template.backends.django.DjangoTemplates',
        'DIRS': [],
    }
]
settings.configure(TEMPLATES=TEMPLATES)
django.setup()
import pdb
from optparse import OptionParser

def docx_contextualize(input_docx,context_file):
    with open(context_file) as csvfile:
        reader = csv.DictReader(csvfile)
        for context in reader:
            print context
            c = Context(context)
            
            f = open(input_docx,'rb')
            doc = Document(f)
            f.close()
            for i in range(0,len(doc.paragraphs)):
                t = Template(doc.paragraphs[i].text)
                doc.paragraphs[i].text = t.render(c)
                
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        t = Template(cell.text)
                        cell.text = t.render(c)
            
            doc.save(context['file_name']+".docx")
            
            
#Parse Options
parser = OptionParser()
parser.add_option("-d", "--docx", dest="docx", default="template.docx",
                help="Input DOCX", metavar="DOCX")
parser.add_option("-c", "--csv", dest="csv", default="context.csv",
                help="Context CSV", metavar="CSV")
(options, args) = parser.parse_args()
docx_contextualize(options.docx,options.csv)